"""Text extraction from PDF / DOCX / XLSX regulatory documents.

Returns plain text plus a page (or sheet) count. PDFs are tried with pypdf
first (fast); if that yields almost nothing we retry with pdfplumber. A near-empty
result flags a likely scanned/image PDF that would need OCR downstream.
"""
from __future__ import annotations

from pathlib import Path
from typing import Tuple

# Below this many characters of extractable text, we treat a PDF as scanned.
SCANNED_CHAR_THRESHOLD = 200


def extract_pdf(path: Path) -> Tuple[str, int]:
    text, pages = _extract_pdf_pypdf(path)
    if len(text.strip()) < SCANNED_CHAR_THRESHOLD:
        # pypdf came up nearly empty — try the heavier, more robust extractor.
        alt_text, alt_pages = _extract_pdf_pdfplumber(path)
        if len(alt_text.strip()) > len(text.strip()):
            return alt_text, alt_pages or pages
    return text, pages


def _extract_pdf_pypdf(path: Path) -> Tuple[str, int]:
    from pypdf import PdfReader

    try:
        reader = PdfReader(str(path))
    except Exception:
        return "", 0
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n\n".join(parts), len(reader.pages)


def _extract_pdf_pdfplumber(path: Path) -> Tuple[str, int]:
    try:
        import pdfplumber
    except Exception:
        return "", 0
    parts = []
    pages = 0
    try:
        with pdfplumber.open(str(path)) as pdf:
            pages = len(pdf.pages)
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
    except Exception:
        return "", pages
    return "\n\n".join(parts), pages


def extract_docx(path: Path) -> Tuple[str, int]:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # include table cell text — common in regulatory checklists
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts), 0


def extract_xlsx(path: Path) -> Tuple[str, int]:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts), len(wb.sheetnames) if hasattr(wb, "sheetnames") else 0


def extract(path: Path) -> Tuple[str, int, bool]:
    """Extract (text, page_count, is_scanned) from a supported document."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        text, pages = extract_pdf(path)
        is_scanned = len(text.strip()) < SCANNED_CHAR_THRESHOLD
        return text, pages, is_scanned
    if ext == ".docx":
        text, pages = extract_docx(path)
        return text, pages, False
    if ext == ".xlsx":
        text, pages = extract_xlsx(path)
        return text, pages, False
    raise ValueError(f"Unsupported file type: {ext}")
